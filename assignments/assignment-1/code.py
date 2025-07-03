# Assignment based on basic RAG introduction
# 1. fetch the data from pdf - system_design.pdf
# 2. if chunking(use the sementic chunking technique) required do chunking and then embedding
# 4. store it inside the vector database(use any of them 1. mongodb 2. astradb 3. opensearch 4.milvus) ## i have not discuss then you need to explore
# Set up Milvus locally: https://milvus.io/docs/install_standalone-windows.md
# 5. create a index with all three index machnism(Flat, HNSW, IVF) ## i have not discuss then you need to explore
# 6. create a retriever pipeline
# 7. check the retriever time(which one is fastet)
# 8. print the accuray score of every similarity search
# 9. perform the reranking either using BM25 or MMR ## i have not discuss then you need to explore
# 10. then write a prompt template
# 11. generte a oputput through llm
# 12. render that output over the DOCx
import os
from dotenv import load_dotenv
from typing import List, Dict, Any
import time
from datetime import datetime

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.schema import Document
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from rank_bm25 import BM25Okapi
from docx import Document as DocxDocument
from tqdm import tqdm
from pymilvus import (
    connections,
    Collection,
    CollectionSchema,
    FieldSchema,
    DataType,
    utility
)

load_dotenv()

# Constants
PDF_PATH = "assignments/assignment-1/system_design.pdf"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
MILVUS_HOST = "localhost"
MILVUS_PORT = "19530"
COLLECTION_PREFIX = "rag_documents"
DIMENSION = 384  # Dimension for all-MiniLM-L6-v2 embeddings

class RAGEngine:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        # Connect to Milvus
        connections.connect(host=MILVUS_HOST, port=MILVUS_PORT)
        
        # Initialize collections for different index types
        self.collections = {}
        self._setup_collections()
        
        self.llm = ChatGroq(
            api_key=GROQ_API_KEY,
            model_name="qwen-qwq-32b"
        )
    
    def _get_collection_schema(self) -> CollectionSchema:
        """Get the schema for collections"""
        fields = [
            FieldSchema(name="id", dtype=DataType.INT64, is_primary=True, auto_id=True),
            FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
            FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=DIMENSION),
            FieldSchema(name="metadata", dtype=DataType.JSON)
        ]
        return CollectionSchema(fields=fields, description="RAG document collection")
    
    def _setup_collections(self):
        """Setup separate collections for each index type"""
        index_types = {
            "flat": {
                "index_params": {
                    "metric_type": "COSINE",
                    "index_type": "FLAT",
                    "params": {}
                }
            },
            "hnsw": {
                "index_params": {
                    "metric_type": "COSINE",
                    "index_type": "HNSW",
                    "params": {
                        "M": 16,
                        "efConstruction": 100
                    }
                }
            },
            "ivf": {
                "index_params": {
                    "metric_type": "COSINE",
                    "index_type": "IVF_FLAT",
                    "params": {
                        "nlist": 100
                    }
                }
            }
        }
        
        schema = self._get_collection_schema()
        
        for index_type, config in index_types.items():
            collection_name = f"{COLLECTION_PREFIX}_{index_type}"
            
            # Drop collection if exists
            if utility.has_collection(collection_name):
                utility.drop_collection(collection_name)
            
            # Create collection
            collection = Collection(name=collection_name, schema=schema)
            
            # Create index
            collection.create_index(
                field_name="embedding",
                index_params=config["index_params"],
                index_name=f"{index_type}_index"
            )
            
            self.collections[index_type] = collection
            print(f"Created collection {collection_name} with {index_type} index")

    def load_and_chunk_documents(self) -> List[Document]:
        """Load PDF and perform chunking"""
        print("Loading and chunking documents...")
        loader = PyPDFLoader(PDF_PATH)
        documents = loader.load()
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        
        chunks = text_splitter.split_documents(documents)
        print(f"Created {len(chunks)} chunks from {len(documents)} pages")
        return chunks

    def create_embeddings_and_store(self, chunks: List[Document]):
        """Create embeddings and store in all Milvus collections"""
        print("Creating embeddings and storing in Milvus collections...")
        
        # Clean and prepare texts for embedding
        texts = []
        metadatas = []
        for chunk in chunks:
            # Clean and validate text
            text = chunk.page_content.strip()
            if text and isinstance(text, str):  # Ensure text is non-empty string
                texts.append(text)
                metadatas.append(chunk.metadata)
        
        if not texts:
            raise ValueError("No valid texts found for embedding")
        
        print(f"Preparing to embed {len(texts)} valid texts...")
        
        # Create embeddings in batches to handle large documents
        batch_size = 32
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            try:
                batch_embeddings = self.embeddings.embed_documents(batch_texts)
                all_embeddings.extend(batch_embeddings)
                print(f"Embedded batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
            except Exception as e:
                print(f"Error embedding batch starting at index {i}: {str(e)}")
                # Try embedding one by one for this batch
                for j, text in enumerate(batch_texts):
                    try:
                        embedding = self.embeddings.embed_query(text)
                        all_embeddings.append(embedding)
                    except Exception as e:
                        print(f"Error embedding text at index {i+j}: {str(e)}")
                        # Use zero vector as fallback
                        all_embeddings.append([0.0] * DIMENSION)
        
        if len(all_embeddings) != len(texts):
            raise ValueError(f"Number of embeddings ({len(all_embeddings)}) doesn't match number of texts ({len(texts)})")
        
        # Insert data into all collections
        entities = [
            texts,  # text field
            all_embeddings,  # embedding field
            metadatas  # metadata field
        ]
        
        for index_type, collection in self.collections.items():
            try:
                collection.insert(entities)
                collection.flush()
                print(f"Stored {len(texts)} documents in {COLLECTION_PREFIX}_{index_type} collection")
            except Exception as e:
                print(f"Error storing in {index_type} collection: {str(e)}")
                # Try inserting one by one
                for i in range(len(texts)):
                    try:
                        single_entity = [
                            [texts[i]],
                            [all_embeddings[i]],
                            [metadatas[i]]
                        ]
                        collection.insert(single_entity)
                    except Exception as e:
                        print(f"Error storing document {i} in {index_type} collection: {str(e)}")
                collection.flush()
                print(f"Completed storing documents in {COLLECTION_PREFIX}_{index_type} collection")

    def evaluate_retrieval_performance(self, query: str, k: int = 5) -> Dict[str, Any]:
        """Evaluate retrieval performance for different index types"""
        results = {}
        
        # Clean and validate query
        query = query.strip()
        if not query or not isinstance(query, str):
            raise ValueError("Invalid query text")
        
        try:
            query_embedding = self.embeddings.embed_query(query)
        except Exception as e:
            print(f"Error creating query embedding: {str(e)}")
            raise
        
        for index_type, collection in self.collections.items():
            try:
                # Load the collection
                collection.load()
                
                # Measure retrieval time
                start_time = time.time()
                
                # Search in Milvus
                search_params = {
                    "metric_type": "COSINE",
                    "params": {"nprobe": 10} if index_type == "ivf" else {}
                }
                
                search_results = collection.search(
                    data=[query_embedding],
                    anns_field="embedding",
                    param=search_params,
                    limit=k,
                    output_fields=["text", "metadata"]
                )
                
                retrieval_time = time.time() - start_time
                
                # Process results
                docs = []
                for hits in search_results:
                    for hit in hits:
                        text = hit.entity.get("text", "").strip()
                        if text:  # Only include non-empty results
                            docs.append(Document(
                                page_content=text,
                                metadata=hit.entity.get("metadata", {})
                            ))
                
                if not docs:
                    print(f"Warning: No results found in {index_type} collection")
                    continue
                
                # Calculate accuracy (using cosine similarity as metric)
                doc_embeddings = []
                for doc in docs:
                    try:
                        doc_embedding = self.embeddings.embed_query(doc.page_content)
                        doc_embeddings.append(doc_embedding)
                    except Exception as e:
                        print(f"Error creating embedding for document in {index_type} collection: {str(e)}")
                        continue
                
                if not doc_embeddings:
                    print(f"Warning: Could not create embeddings for any documents in {index_type} collection")
                    continue
                
                similarities = cosine_similarity([query_embedding], doc_embeddings)[0]
                avg_similarity = np.mean(similarities)
                
                results[index_type] = {
                    "retrieval_time": retrieval_time,
                    "avg_similarity": avg_similarity,
                    "documents": docs
                }
                
                print(f"\n{index_type.upper()} Index Results:")
                print(f"Collection: {COLLECTION_PREFIX}_{index_type}")
                print(f"Retrieval Time: {retrieval_time:.4f} seconds")
                print(f"Average Similarity: {avg_similarity:.4f}")
                print(f"Number of documents retrieved: {len(docs)}")
                
            except Exception as e:
                print(f"Error evaluating {index_type} collection: {str(e)}")
                continue
            finally:
                # Release the collection
                try:
                    collection.release()
                except Exception as e:
                    print(f"Error releasing {index_type} collection: {str(e)}")
        
        if not results:
            raise ValueError("No successful retrievals from any collection")
            
        return results

    def rerank_documents(self, docs: List[Document], query: str, method: str = "bm25") -> List[Document]:
        """Rerank documents using either BM25 or MMR"""
        if method == "bm25":
            # BM25 reranking
            tokenized_docs = [doc.page_content.split() for doc in docs]
            bm25 = BM25Okapi(tokenized_docs)
            tokenized_query = query.split()
            scores = bm25.get_scores(tokenized_query)
            
            # Sort documents by BM25 scores
            doc_score_pairs = list(zip(docs, scores))
            reranked_docs = [doc for doc, _ in sorted(doc_score_pairs, key=lambda x: x[1], reverse=True)]
            
        else:  # MMR
            # Implement MMR reranking
            query_embedding = self.embeddings.embed_query(query)
            doc_embeddings = [self.embeddings.embed_query(doc.page_content) for doc in docs]
            
            # Calculate relevance and diversity scores
            relevance_scores = [cosine_similarity([query_embedding], [doc_emb])[0][0] for doc_emb in doc_embeddings]
            reranked_docs = []
            remaining_docs = list(zip(docs, doc_embeddings, relevance_scores))
            
            while remaining_docs:
                # Select document with highest relevance
                selected_idx = np.argmax([score for _, _, score in remaining_docs])
                selected_doc, selected_emb, _ = remaining_docs.pop(selected_idx)
                reranked_docs.append(selected_doc)
                
                # Update relevance scores for remaining documents
                for i, (_, doc_emb, _) in enumerate(remaining_docs):
                    diversity_score = cosine_similarity([selected_emb], [doc_emb])[0][0]
                    remaining_docs[i] = (remaining_docs[i][0], doc_emb, 
                                       remaining_docs[i][2] - 0.5 * diversity_score)
        
        return reranked_docs

    def generate_response(self, query: str, context: List[Document]) -> str:
        """Generate response using Groq LLM"""
        prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""You are a helpful AI assistant. Use the following context to answer the question.
            If you don't know the answer, just say that you don't know. Don't try to make up an answer.
            
            Context: {context}
            
            Question: {question}
            
            Answer: """
        )
        
        context_str = "\n\n".join([doc.page_content for doc in context])
        prompt = prompt_template.format(context=context_str, question=query)
        
        response = self.llm.invoke(prompt)
        return response.content

    def save_to_docx(self, content: str, output_path: str):
        """Save the generated content to a DOCX file"""
        doc = DocxDocument()
        doc.add_paragraph(content)
        doc.save(output_path)

def main():
    # Initialize RAG engine
    rag = RAGEngine()
    
    # Load and process documents
    chunks = rag.load_and_chunk_documents()
    rag.create_embeddings_and_store(chunks)
    
    # Example query
    query = "What are the key components of system design?"
    
    # Evaluate retrieval performance
    print("\nEvaluating retrieval performance...")
    results = rag.evaluate_retrieval_performance(query)
    
    # Select best performing index (based on retrieval time)
    best_index = min(results.items(), key=lambda x: x[1]['retrieval_time'])[0]
    print(f"\nBest performing index: {best_index}")
    print(f"Best performing collection: {COLLECTION_PREFIX}_{best_index}")
    
    # Rerank documents
    print("\nReranking documents...")
    reranked_docs = rag.rerank_documents(results[best_index]['documents'], query, method="bm25")
    
    # Generate response
    print("\nGenerating response...")
    response = rag.generate_response(query, reranked_docs)
    
    # Save to DOCX
    output_path = "assignments/assignment-1/rag_response.docx"
    rag.save_to_docx(response, output_path)
    print(f"\nResponse saved to {output_path}")

if __name__ == "__main__":
    main()